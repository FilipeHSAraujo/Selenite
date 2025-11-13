# utils/reporter.py
"""
SELENITE V0.1 – Test Reporter Module
Generates:
- Excel: Full test plan execution
- Word: Only defects (ready for Jira/Azure/Dev)
Author: Filipe Araujo
"""

import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches
from colorama import Fore, Style, init

init(autoreset=True)


class TestReporter:
    def __init__(self, plan_name: str, plan_url: str = ""):
        """
        Initialize reporter with plan name and URL.
        Creates reports/ and screenshots/ folders automatically.
        """
        self.plan_name = plan_name.replace(" ", "_").replace("/", "_")
        self.plan_url = plan_url or "URL not provided"
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.results = []
        self.report_dir = f"reports/{self.plan_name}_{self.timestamp}"
        os.makedirs(self.report_dir, exist_ok=True)
        os.makedirs(f"{self.report_dir}/screenshots", exist_ok=True)

    def add_step(self, step_desc: str, status: bool, duration=None, error=None, screenshot_path=None):
        """
        Add a test step result.
        """
        result = {
            "Step": step_desc,
            "Status": "PASS" if status else "FAIL",
            "Duration (s)": round(duration, 3) if duration else "-",
            "Error": str(error) if error else "-",
            "Screenshot": os.path.basename(screenshot_path) if screenshot_path else "-"
        }
        self.results.append(result)

    def save_screenshot(self, driver, name: str) -> str:
        """
        Save screenshot with safe filename.
        Returns full path.
        """
        safe_name = "".join(c for c in name if c.isalnum() or c in (" ", "_", "-")).strip()
        safe_name = safe_name.replace(" ", "_")
        path = f"{self.report_dir}/screenshots/{safe_name}.png"
        driver.save_screenshot(path)
        return path

    def generate_excel(self):
        """
        Generate full execution plan in Excel (.xlsx)
        """
        df = pd.DataFrame(self.results)
        excel_path = f"{self.report_dir}/{self.plan_name}_EXECUTION_PLAN.xlsx"
        df.to_excel(excel_path, index=False)
        print(f"{Fore.GREEN}EXECUTION PLAN → {excel_path}{Style.RESET_ALL}")

    def generate_word_defects(self):
        """
        Generate Word report ONLY with defects.
        Includes: steps to reproduce, error, full HD screenshot.
        """
        if not any(r["Status"] == "FAIL" for r in self.results):
            print(f"{Fore.YELLOW}No failures! Defect report skipped.{Style.RESET_ALL}")
            return

        doc = Document()
        doc.add_heading("DEFECT REPORT - SELENITE", 0)
        doc.add_paragraph(f"Project: {self.plan_name}")
        doc.add_paragraph(f"URL: {self.plan_url}")
        doc.add_paragraph(f"Executed on: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        doc.add_paragraph("")

        bug_id = 0
        for i, res in enumerate(self.results):
            if res["Status"] == "FAIL":
                bug_id += 1
                doc.add_heading(f"BUG {bug_id} - {res['Step']}", level=1)

                # Steps to reproduce
                p = doc.add_paragraph()
                p.add_run("Steps to reproduce:\n").bold = True
                for j, step in enumerate(self.results[:i + 1]):
                    status_icon = "PASS" if step["Status"] == "PASS" else "FAIL"
                    p.add_run(f"{j + 1}. {step['Step']} [{status_icon}]\n")

                # Technical error
                if res["Error"] != "-":
                    doc.add_paragraph(f"Error: {res['Error']}", style="Intense Quote")

                # Evidence
                doc.add_paragraph("Evidence:", style="Intense Quote")
                if res["Screenshot"] != "-":
                    img_path = f"{self.report_dir}/screenshots/{res['Screenshot']}"
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(6.8))
                        doc.add_paragraph(f"Screenshot: {res['Screenshot']}")

                doc.add_page_break()

        word_path = f"{self.report_dir}/{self.plan_name}_DEFECTS.docx"
        doc.save(word_path)
        print(f"{Fore.RED}DEFECTS REPORT → {word_path}{Style.RESET_ALL}")

    def generate_reports(self):
        """
        Generate both reports:
        - Excel: full execution
        - Word: only defects
        """
        self.generate_excel()
        self.generate_word_defects()
        print(f"{Fore.MAGENTA}ALL REPORTS READY → {self.report_dir}{Style.RESET_ALL}")